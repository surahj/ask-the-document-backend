"""
Document processing service for DocuMind AI Assistant with vector embeddings
"""

import os
import re
import tempfile
import requests
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import markdown
from sqlalchemy.orm import Session
from app.config import settings
from app.models import Document, DocumentChunk
from app.services.embedding_service import EmbeddingService
from app.services.cloudinary_service import CloudinaryService


class DocumentProcessor:
    """Document processing service with vector embedding support"""

    def __init__(self, embedding_service: EmbeddingService = None):
        self.supported_extensions = settings.allowed_extensions
        self.max_file_size = settings.max_file_size
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.max_chunks = settings.max_chunks_per_document
        self.embedding_service = embedding_service or EmbeddingService()
        self.cloudinary_service = CloudinaryService()

    def validate_document(self, file_path: str) -> Dict[str, Any]:
        """Validate document file"""
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                return {"error": "File not found"}

            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size > self.max_file_size:
                return {
                    "error": f"File size exceeds maximum limit of {self.max_file_size} bytes"
                }

            # Check file extension
            file_ext = Path(file_path).suffix.lower()
            if file_ext not in self.supported_extensions:
                return {
                    "error": f"Unsupported file format. Supported formats: {', '.join(self.supported_extensions)}"
                }

            return {"valid": True, "file_size": file_size, "file_type": file_ext}

        except Exception as e:
            return {"error": f"Validation error: {str(e)}"}

    def download_from_cloudinary(
        self, cloudinary_url: str, file_type: str
    ) -> Dict[str, Any]:
        """Download file from Cloudinary URL for processing"""
        try:
            # Add timeout and better error handling
            response = requests.get(cloudinary_url, stream=True, timeout=30)
            response.raise_for_status()

            # Create temporary file
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=file_type
            ) as temp_file:
                for chunk in response.iter_content(chunk_size=8192):
                    temp_file.write(chunk)
                temp_file_path = temp_file.name

            return {"success": True, "temp_file_path": temp_file_path}

        except requests.exceptions.HTTPError as e:
            if e.response.status_code == 401:
                return {
                    "error": "Cloudinary authentication failed. Please check your Cloudinary credentials in the .env file."
                }
            elif e.response.status_code == 404:
                return {
                    "error": "File not found in Cloudinary. The file may have been deleted or moved."
                }
            elif e.response.status_code == 403:
                return {
                    "error": "Access denied to Cloudinary file. The file may be private or require different permissions."
                }
            else:
                return {
                    "error": f"Failed to download from Cloudinary: HTTP {e.response.status_code} - {e.response.text}"
                }
        except requests.exceptions.Timeout:
            return {
                "error": "Timeout while downloading from Cloudinary. Please try again."
            }
        except requests.exceptions.ConnectionError:
            return {
                "error": "Connection error while accessing Cloudinary. Please check your internet connection."
            }
        except Exception as e:
            return {"error": f"Failed to download from Cloudinary: {str(e)}"}

    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from document"""
        try:
            file_ext = Path(file_path).suffix.lower()

            if file_ext == ".pdf":
                return self._extract_pdf_text(file_path)
            elif file_ext == ".docx":
                return self._extract_docx_text(file_path)
            elif file_ext in [".txt", ".md"]:
                return self._extract_text_file(file_path)
            else:
                return {"error": f"Unsupported file format: {file_ext}"}

        except Exception as e:
            return {"error": f"Text extraction error: {str(e)}"}

    def extract_text_from_cloudinary(
        self, cloudinary_url: str, file_type: str
    ) -> Dict[str, Any]:
        """Extract text from document stored in Cloudinary"""
        try:
            # Download file from Cloudinary
            download_result = self.download_from_cloudinary(cloudinary_url, file_type)
            if "error" in download_result:
                return download_result

            temp_file_path = download_result["temp_file_path"]

            try:
                # Extract text from the temporary file
                extraction_result = self.extract_text(temp_file_path)
                return extraction_result
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)

        except Exception as e:
            return {"error": f"Text extraction from Cloudinary error: {str(e)}"}

    def _extract_pdf_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""

                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                return {"text": text, "pages": len(pdf_reader.pages)}

        except Exception as e:
            return {"error": f"PDF extraction error: {str(e)}"}

    def _extract_docx_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            return {"text": text, "paragraphs": len(doc.paragraphs)}

        except Exception as e:
            return {"error": f"DOCX extraction error: {str(e)}"}

    def _extract_text_file(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text or markdown file"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()

            # Convert markdown to plain text if needed
            if file_path.endswith(".md"):
                content = markdown.markdown(content)
                # Remove HTML tags
                content = re.sub(r"<[^>]+>", "", content)

            return {"text": content, "lines": len(content.split("\n"))}

        except Exception as e:
            return {"error": f"Text file extraction error: {str(e)}"}

    def chunk_text(self, text: str) -> List[Dict[str, Any]]:
        """Chunk text into smaller pieces"""
        try:
            if not text.strip():
                return []

            # Clean and normalize text
            text = self._clean_text(text)

            # Split into sentences first
            sentences = self._split_into_sentences(text)

            chunks = []
            current_chunk = ""
            current_start = 0
            chunk_index = 0

            for i, sentence in enumerate(sentences):
                # Check if adding this sentence would exceed chunk size
                if (
                    len(current_chunk) + len(sentence) > self.chunk_size
                    and current_chunk
                ):
                    # Save current chunk
                    chunks.append(
                        {
                            "chunk_index": chunk_index,
                            "content": current_chunk.strip(),
                            "start_position": current_start,
                            "end_position": current_start + len(current_chunk),
                        }
                    )

                    # Start new chunk with overlap
                    overlap_text = self._get_overlap_text(current_chunk)
                    current_chunk = overlap_text + sentence
                    current_start = (
                        current_start
                        + len(current_chunk)
                        - len(overlap_text)
                        - len(sentence)
                    )
                    chunk_index += 1

                    # Check max chunks limit
                    if chunk_index >= self.max_chunks:
                        break
                else:
                    current_chunk += sentence + " "

            # Add final chunk if there's content
            if current_chunk.strip() and chunk_index < self.max_chunks:
                chunks.append(
                    {
                        "chunk_index": chunk_index,
                        "content": current_chunk.strip(),
                        "start_position": current_start,
                        "end_position": current_start + len(current_chunk),
                    }
                )

            return chunks

        except Exception as e:
            return [{"error": f"Chunking error: {str(e)}"}]

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r"\s+", " ", text)
        # Remove special characters but keep punctuation
        text = re.sub(r"[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}]", "", text)
        return text.strip()

    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting - can be improved with NLP libraries
        sentences = re.split(r"[.!?]+", text)
        return [s.strip() for s in sentences if s.strip()]

    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of a chunk"""
        if len(text) <= self.chunk_overlap:
            return text
        return text[-self.chunk_overlap :]

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process document and return chunks with metadata"""
        try:
            # Validate document
            validation = self.validate_document(file_path)
            if "error" in validation:
                return validation

            # Extract text
            extraction = self.extract_text(file_path)
            if "error" in extraction:
                return extraction

            # Chunk text
            chunks = self.chunk_text(extraction["text"])
            if not chunks or "error" in chunks[0]:
                return {"error": "Failed to chunk document"}

            return {
                "chunks": chunks,
                "file_info": {
                    "file_size": validation["file_size"],
                    "file_type": validation["file_type"],
                    "total_chunks": len(chunks),
                },
            }

        except Exception as e:
            return {"error": f"Document processing error: {str(e)}"}

    def process_document_bytes(
        self,
        file_content: bytes,
        file_ext: str,
        user_id: int,
        filename: str,
        db: Session,
    ) -> Dict[str, Any]:
        """Process document from bytes (no Cloudinary download)"""
        try:
            # Save to a temporary file for processing
            import tempfile

            with tempfile.NamedTemporaryFile(
                delete=False, suffix=file_ext
            ) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            try:
                # Validate document
                validation = self.validate_document(temp_file_path)
                if "error" in validation:
                    return validation
                # Extract text
                extraction = self.extract_text(temp_file_path)
                if "error" in extraction:
                    return extraction
                # Chunk text
                chunks = self.chunk_text(extraction["text"])
                if not chunks or (isinstance(chunks, list) and "error" in chunks[0]):
                    return {"error": "Failed to chunk document"}
                return {"chunks": chunks}
            finally:
                import os

                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
        except Exception as e:
            return {"error": f"Document processing error: {str(e)}"}

    def save_document_record(
        self,
        user_id: int,
        filename: str,
        file_size: int,
        file_type: str,
        db: Session,
        cloudinary_url: Optional[str] = None,
        cloudinary_public_id: Optional[str] = None,
        chunks: Optional[List[Dict[str, Any]]] = None,
    ) -> int:
        """Save document and chunks to DB, including Cloudinary URL if provided"""
        from app.models import Document, DocumentChunk

        try:
            document = Document(
                user_id=user_id,
                filename=filename,
                file_size=file_size,
                file_type=file_type,
                status="processed",
                cloudinary_url=cloudinary_url,
                cloudinary_public_id=cloudinary_public_id,
            )
            db.add(document)
            db.flush()  # Get document.id
            # Save chunks
            if chunks:
                for chunk_data in chunks:
                    chunk = DocumentChunk(
                        document_id=document.id,
                        chunk_index=chunk_data["chunk_index"],
                        content=chunk_data["content"],
                        start_position=chunk_data["start_position"],
                        end_position=chunk_data["end_position"],
                    )
                    db.add(chunk)
            db.commit()
            return document.id
        except Exception as e:
            db.rollback()
            import logging

            logging.error(f"Failed to save document record: {e}")
            raise

    def process_and_store_document(
        self, file_path: str, user_id: int, filename: str, db: Session
    ) -> Dict[str, Any]:
        """Process document and store with embeddings in database"""
        try:
            # Check if document with same filename already exists for this user
            existing_doc = (
                db.query(Document)
                .filter(Document.user_id == user_id, Document.filename == filename)
                .first()
            )

            if existing_doc:
                return {
                    "error": f"Document '{filename}' already exists. Please use a different filename or delete the existing document first."
                }

            # Process document
            result = self.process_document(file_path)
            if "error" in result:
                return result

            # Create document record
            document = Document(
                user_id=user_id,
                filename=filename,
                file_path=file_path,
                file_size=result["file_info"]["file_size"],
                file_type=result["file_info"]["file_type"],
                status="processing",
            )
            db.add(document)
            db.flush()  # Get the document ID

            # Create chunks and embeddings
            chunks_created = 0
            for chunk_data in result["chunks"]:
                # Create chunk record
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=chunk_data["chunk_index"],
                    content=chunk_data["content"],
                    start_position=chunk_data["start_position"],
                    end_position=chunk_data["end_position"],
                )
                db.add(chunk)
                db.flush()  # Get the chunk ID

                # Create and store embedding
                embedding = self.embedding_service.create_embedding(
                    chunk_data["content"]
                )
                if embedding:
                    # Store embedding directly as list for PostgreSQL VECTOR type
                    chunk.embedding = embedding
                    chunks_created += 1
                else:
                    print(
                        f"Warning: Failed to create embedding for chunk {chunk_data['chunk_index']}"
                    )

            # Update document status
            document.status = "processed"
            db.commit()

            return {
                "document_id": document.id,
                "chunks_created": chunks_created,
                "total_chunks": len(result["chunks"]),
                "file_info": result["file_info"],
                "status": "processed",
            }

        except Exception as e:
            db.rollback()
            return {"error": f"Error processing and storing document: {str(e)}"}

    def process_and_store_document_from_cloudinary(
        self,
        cloudinary_url: str,
        cloudinary_public_id: str,
        user_id: int,
        filename: str,
        file_size: int,
        file_type: str,
        db: Session,
    ) -> Dict[str, Any]:
        """Process document from Cloudinary URL and store with embeddings in database"""
        try:
            # Check if document with same filename already exists for this user
            existing_doc = (
                db.query(Document)
                .filter(Document.user_id == user_id, Document.filename == filename)
                .first()
            )

            if existing_doc:
                return {
                    "error": f"Document '{filename}' already exists. Please use a different filename or delete the existing document first."
                }

            # Extract text from Cloudinary
            extraction = self.extract_text_from_cloudinary(cloudinary_url, file_type)
            if "error" in extraction:
                return extraction

            # Chunk text
            chunks = self.chunk_text(extraction["text"])
            if not chunks or "error" in chunks[0]:
                return {"error": "Failed to chunk document"}

            # Create document record
            document = Document(
                user_id=user_id,
                filename=filename,
                cloudinary_url=cloudinary_url,
                cloudinary_public_id=cloudinary_public_id,
                file_size=file_size,
                file_type=file_type,
                status="processing",
            )
            db.add(document)
            db.flush()  # Get the document ID

            # Create chunks and embeddings
            chunks_created = 0
            for chunk_data in chunks:
                # Create chunk record
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=chunk_data["chunk_index"],
                    content=chunk_data["content"],
                    start_position=chunk_data["start_position"],
                    end_position=chunk_data["end_position"],
                )
                db.add(chunk)
                db.flush()  # Get the chunk ID

                # Create and store embedding
                embedding = self.embedding_service.create_embedding(
                    chunk_data["content"]
                )
                if embedding:
                    # Store embedding directly as list for PostgreSQL VECTOR type
                    chunk.embedding = embedding
                    chunks_created += 1
                else:
                    print(
                        f"Warning: Failed to create embedding for chunk {chunk_data['chunk_index']}"
                    )

            # Update document status
            document.status = "processed"
            db.commit()

            return {
                "document_id": document.id,
                "chunks_created": chunks_created,
                "total_chunks": len(chunks),
                "file_info": {
                    "file_size": file_size,
                    "file_type": file_type,
                    "total_chunks": len(chunks),
                },
                "status": "processed",
            }

        except Exception as e:
            db.rollback()
            return {
                "error": f"Error processing and storing document from Cloudinary: {str(e)}"
            }

    def delete_document_chunks(self, document_id: int, db: Session) -> bool:
        """Delete all chunks and embeddings for a document"""
        try:
            # Delete chunks (embeddings will be deleted due to cascade)
            chunks = (
                db.query(DocumentChunk)
                .filter(DocumentChunk.document_id == document_id)
                .all()
            )

            for chunk in chunks:
                db.delete(chunk)

            db.commit()
            return True

        except Exception as e:
            db.rollback()
            print(f"Error deleting document chunks: {e}")
            return False
